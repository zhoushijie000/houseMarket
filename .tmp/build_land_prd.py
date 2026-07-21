from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\ZhuanZ\Desktop\zhou\houseMarket")
ASSETS = ROOT / ".tmp" / "prd-assets"
OUTPUT = ROOT / "output" / "成都住建房产超市_土拍分析需求文档_V1.1_20260720.docx"
TABLE_HELPER_DIR = Path(
    r"C:\Users\ZhuanZ\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents\scripts"
)
sys.path.insert(0, str(TABLE_HELPER_DIR))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


FONT_BODY = "Microsoft YaHei"
FONT_HEADING = "Microsoft YaHei"
COLOR_TEXT = "222222"
COLOR_MUTED = "666666"
COLOR_BLUE = "2E74B5"
COLOR_ORANGE = "FF6A21"
COLOR_HEADER = "F2F2F2"
COLOR_BORDER = "D9D9D9"
TABLE_WIDTH = 9360


def set_run_font(run, name=FONT_BODY, size=10.5, bold=None, color=COLOR_TEXT, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=COLOR_BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_horizontal_rule(paragraph, color="D9D9D9", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    prefix = p.add_run("第 ")
    set_run_font(prefix, size=8.5, color="7F7F7F")
    page_run = p.add_run()
    set_run_font(page_run, size=8.5, color="7F7F7F")
    add_field(page_run, "PAGE")
    suffix = p.add_run(" 页")
    set_run_font(suffix, size=8.5, color="7F7F7F")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_specs = {
        "Heading 1": (15, 16, 8),
        "Heading 2": (12.5, 12, 6),
        "Heading 3": (11.5, 9, 4),
    }
    for style_name, (size, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = FONT_HEADING
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEADING)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEADING)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167

    add_footer(section)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("成都住建房产超市土拍分析")
    set_run_font(run, name=FONT_HEADING, size=18, bold=True, color="1F1F1F")

    metadata = [
        ("升级时间", "2026/7/20"),
        ("版本", "V1.1"),
        ("产品", "周世杰"),
        ("评审时间", "/"),
        ("系统", "成都住建房产超市"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        label_run = p.add_run(f"{label}：")
        set_run_font(label_run, size=9.5, color=COLOR_MUTED)
        value_run = p.add_run(value)
        set_run_font(value_run, size=9.5, color=COLOR_MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    add_horizontal_rule(rule)


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def add_paragraph(doc, text="", bold_prefix=None, color=COLOR_TEXT, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc, headers, rows, weights=None, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header_row = table.rows[0]
    set_repeat_table_header(header_row)

    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, COLOR_HEADER)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)

    widths = column_widths_from_weights(weights or [1] * len(headers), TABLE_WIDTH)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=TABLE_WIDTH,
        indent_dxa=110,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 110, "end": 110},
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_screenshot(doc, filename, caption, width=3.15, page_break=False):
    if page_break:
        doc.add_page_break()
    path = ASSETS / filename
    if not path.exists():
        p = doc.add_paragraph(f"[截图缺失：{filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=8.5, color=COLOR_MUTED)


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF4ED")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "5")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "FFD7BF")
        borders.append(border)
    p_pr.append(borders)
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, size=9.5, bold=True, color=COLOR_ORANGE)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.5)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "1. 功能清单", 1)
    add_heading(doc, "1.1 C端功能清单", 2)
    add_table(
        doc,
        ["一级功能", "二级功能", "功能描述（调整内容）", "需求类型", "后端研发", "前端研发", "上线时间"],
        [
            ["首页", "土拍分析入口", "新增土拍分析入口，点击进入土拍分析页", "新增", "待定", "待定", "待定"],
            ["土拍分析", "成交榜单", "支持成交楼面价榜、土地溢价率榜及完整榜单", "新增", "待定", "待定", "待定"],
            ["土拍分析", "地块检索", "支持关键词、状态、区域、用途和排序组合筛选", "新增", "待定", "待定", "待定"],
            ["土拍分析", "地块列表", "展示状态、名称、起拍楼面价、成交价、发布时间、成交时间、区域、面积、容积率", "新增", "待定", "待定", "待定"],
            ["土拍分析", "土拍详情", "展示宗地基础信息、成交信息及数据免责声明", "新增", "待定", "待定", "待定"],
        ],
        weights=[1.0, 1.15, 2.3, 0.8, 0.85, 0.85, 0.9],
        font_size=8.2,
    )
    add_heading(doc, "1.2 B端功能清单", 2)
    add_table(
        doc,
        ["一级功能", "二级功能", "功能描述（调整内容）", "需求类型", "备注"],
        [["土拍管理", "数据维护", "本次未提供B端原型与管理规则，暂不纳入本版本开发范围", "待确认", "如需上线真实数据，需补充采集、审核、发布、纠错与权限方案"]],
        weights=[1.1, 1.2, 3.0, 1.0, 2.3],
        font_size=8.6,
    )
    add_paragraph(doc, "原型文件：土拍分析.html；首页入口：安家成都-首页-扫码.html", bold_prefix="原型文件：")
    add_paragraph(doc, "需求依据：当前本地静态原型、页面源码及《成都住建房产超市促销二维码》需求文档版式。", color=COLOR_MUTED, italic=True)

    add_heading(doc, "2. 功能描述", 1, page_break=True)
    add_heading(doc, "2.1 功能概述", 2)
    add_paragraph(doc, "面向购房用户提供成都土拍信息的查询、筛选、排行和详情浏览能力，帮助用户快速了解地块出让状态、成交价格、溢价率和规划指标。")
    add_bullets(
        doc,
        [
            "用户入口：首页功能卡片“土拍分析”。",
            "核心链路：首页入口 -> 土拍分析列表 -> 组合筛选/榜单 -> 地块详情。",
            "数据范围：成都全市公开土拍信息；原型数据为静态示例，正式数据源与更新频率待确认。",
            "适用终端：房产超市移动端H5/小程序WebView，文档按移动端交互描述。",
        ],
    )
    add_callout(doc, "范围说明", "本版本聚焦C端查询展示；B端数据运营、人工审核和发布流程不在当前范围内。")

    add_heading(doc, "2.2 首页入口", 2)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["入口位置", "首页功能卡片“土拍分析”"],
            ["点击结果", "进入土拍分析首页，默认展示成交楼面价榜、全部状态和最新发布排序"],
            ["返回规则", "土拍分析页左上角返回首页；详情/榜单详情返回时恢复进入前的列表滚动位置"],
            ["权限", "公开浏览，原型未设置登录或实名校验"],
        ],
        weights=[1.5, 5.0],
    )

    add_heading(doc, "2.3 土拍分析首页", 2)
    add_screenshot(doc, "land-list.png", "图1 土拍分析首页（移动端原型）", width=3.15)
    add_heading(doc, "2.3.1 页面模块与交互", 3)
    add_table(
        doc,
        ["模块/元素", "功能说明", "交互与结果", "状态规则"],
        [
            ["顶部导航", "返回、标题、更多操作", "返回首页；更多操作能力待确认", "返回按钮始终可用"],
            ["搜索框", "按地块名、编号、竞得公司搜索", "输入即刷新列表；支持清空", "无匹配结果显示空状态"],
            ["成交榜单", "展示成交楼面价榜/土拍溢价率榜TOP3", "切换榜单；点击榜单行或完整榜单进入榜单详情", "默认成交楼面价榜"],
            ["状态Tab", "全部、待出让、成交、流拍、终止及对应数量", "单选并即时过滤列表", "默认全部；当前项高亮"],
            ["筛选器", "区域、用途、排序", "打开底部筛选面板；选择后确认生效", "已选条件在筛选按钮上回显"],
            ["地块卡片", "展示地块摘要信息；成交地块额外展示成交价", "点击或键盘Enter/空格进入详情", "非成交地块不展示成交价，成交时间固定显示“-”"],
            ["免责声明", "说明数据来源与使用边界", "仅展示", "列表及详情/榜单区域底部可见"],
        ],
        weights=[1.2, 2.2, 2.3, 1.8],
        font_size=8.6,
    )

    add_heading(doc, "2.3.2 搜索与筛选规则", 3)
    add_table(
        doc,
        ["条件", "类型", "选项/口径", "默认值", "规则"],
        [
            ["关键词", "文本", "地块名、地块编号、竞得公司、所属板块", "空", "去除首尾空格；支持包含匹配；输入即筛选"],
            ["状态", "单选", "全部、待出让、成交、流拍、终止", "全部", "与其他条件取交集"],
            ["区域", "单选", "不限区域及系统返回的行政区", "不限区域", "确认后生效；重置恢复不限"],
            ["用途", "单选", "不限用途、住宅、商住、商办等", "不限用途", "选项由数据字典返回"],
            ["排序", "单选", "最新发布、楼面价从高到低、地块面积从大到小", "最新发布", "最新发布按发布时间倒序"],
        ],
        weights=[1.0, 0.8, 2.6, 1.1, 2.2],
        font_size=8.6,
    )
    add_callout(doc, "原型差异", "搜索框提示包含“编号”，但当前源码检索字段未包含记录ID/土地宗号；正式实现需补齐编号检索。")

    add_heading(doc, "2.3.3 地块列表字段", 3)
    add_table(
        doc,
        ["字段", "定义", "类型/格式", "展示规则"],
        [
            ["地块状态", "当前出让状态", "枚举", "待出让/成交/流拍/终止；不同颜色标签"],
            ["地块名称", "用于识别宗地的标题", "文本", "最多2行，超出省略"],
            ["起拍楼面价", "地块起始楼面单价", "整数，元/㎡", "所有地块展示，千分位格式"],
            ["成交价", "地块成交楼面单价", "整数，元/㎡", "仅成交状态展示，重点色突出"],
            ["区域", "行政区-板块", "文本", "示例：金牛区-国宾"],
            ["发布时间", "土拍信息首次对外发布时间", "YYYY-MM-DD", "所有地块展示"],
            ["成交时间", "土地实际成交日期", "YYYY-MM-DD", "成交地块展示日期；未成交统一显示“-”"],
            ["面积", "净用地面积", "数值，亩", "保留数据源精度"],
            ["容积率", "最大容积率", "数值", "保留数据源精度"],
        ],
        weights=[1.0, 2.2, 1.5, 2.6],
        font_size=8.7,
    )

    add_heading(doc, "2.4 筛选面板", 2, page_break=True)
    add_screenshot(doc, "filter-sheet.png", "图2 区域筛选面板", width=3.15)
    add_bullets(
        doc,
        [
            "点击区域、用途或排序后，以底部弹层展示对应单选项。",
            "弹层打开时锁定页面滚动，顶部筛选按钮显示展开态。",
            "选择选项仅更新草稿值；点击“确定”后应用并刷新列表。",
            "点击“重置”恢复当前筛选维度的默认值；点击关闭、蒙层或Esc键不保存草稿值。",
            "关闭后焦点返回触发筛选的按钮，满足键盘可访问性。",
        ],
    )

    add_heading(doc, "2.5 榜单详情", 2)
    add_screenshot(doc, "ranking-detail.png", "图3 历史成交楼面价完整榜单", width=3.15)
    add_heading(doc, "2.5.1 榜单规则", 3)
    add_table(
        doc,
        ["项目", "成交楼面价榜", "土地溢价率榜"],
        [
            ["统计对象", "成都全市历史成交地块", "已成交且有溢价率的地块"],
            ["排序", "成交楼面价从高到低", "溢价率从高到低"],
            ["首页展示", "TOP3", "TOP3"],
            ["详情展示", "TOP10及统计周期/更新时间", "TOP10及统计周期/更新时间"],
            ["卡片字段", "排名、地块名、区域、竞得方、日期、面积、成交楼面价", "排名、地块名、区域、竞得方、成交楼面价、溢价率"],
        ],
        weights=[1.3, 2.8, 2.8],
        font_size=8.7,
    )
    add_bullets(
        doc,
        [
            "默认与列表页当前榜单类型保持一致。",
            "在榜单详情内切换榜单时，标题、统计摘要、完整榜单同步更新。",
            "榜单更新时间展示最近一次数据完成入库/审核的日期，而非客户端当前日期。",
            "榜单数量不足10条时按实际数量展示，不补空位。",
        ],
    )

    add_heading(doc, "2.6 土拍详情", 2, page_break=True)
    add_screenshot(doc, "land-detail-full.png", "图4 地块详情（成交信息字段）", width=3.15)
    add_heading(doc, "2.6.1 页面模块", 3)
    add_table(
        doc,
        ["模块", "说明", "交互"],
        [
            ["顶部导航", "返回、页面标题、分享", "返回列表并恢复滚动位置；分享能力待确认"],
            ["地块标题", "展示当前地块名称", "仅展示"],
            ["成交摘要", "成交楼面价、成交总价、溢价率", "无值时显示“-”"],
            ["基础信息", "土地宗号、位置、面积、容积率、用途与规划指标", "仅展示"],
            ["成交信息", "交易状态、发布时间、成交时间、竞得方、成交价格与溢价率", "仅展示"],
        ],
        weights=[1.4, 3.4, 2.2],
        font_size=8.8,
    )

    add_heading(doc, "2.6.2 基础信息字段", 3)
    add_table(
        doc,
        ["字段", "定义", "类型/格式", "空值规则"],
        [
            ["土地宗号", "政府公示宗地编号", "文本", "显示“-”"],
            ["宗地位置", "地块详细位置", "文本", "显示“-”"],
            ["净用地面积", "宗地净用地面积", "数值，亩", "显示“-”"],
            ["最大容积率", "规划最大容积率", "数值", "显示“-”"],
            ["用途年限", "土地使用权年限", "整数，年", "显示“-”"],
            ["用地使用性质", "规划用地性质全称", "文本", "显示“-”"],
            ["土地用途", "住宅/商住/商办等", "枚举/文本", "显示“-”"],
            ["绿地率", "规划绿地率约束", "文本/百分比", "显示“-”"],
            ["建筑密度", "规划建筑密度约束", "文本/百分比", "显示“-”"],
            ["建筑高度", "规划建筑限高", "文本", "显示“-”"],
        ],
        weights=[1.25, 2.6, 1.5, 1.8],
        font_size=8.7,
    )

    add_heading(doc, "2.6.3 成交信息字段", 3)
    add_table(
        doc,
        ["字段", "定义", "类型/格式", "展示规则"],
        [
            ["交易状况", "已成交/待出让/已流拍/已终止", "枚举", "由状态映射"],
            ["发布时间", "土拍信息首次对外发布时间", "YYYY-MM-DD", "所有状态展示"],
            ["成交时间", "土地实际成交日期", "YYYY-MM-DD", "未成交统一显示“-”"],
            ["竞得方", "成交受让主体", "文本", "非成交状态显示“-”"],
            ["成交总价", "宗地成交总价", "万元，2位小数", "重点色展示"],
            ["成交价", "按亩计算成交价", "万元/亩，4位小数", "无值显示“-”"],
            ["成交楼面地价", "按计容建筑面积计算成交价", "元/平方米，整数", "千分位展示"],
            ["溢价率", "相对起始价的溢价比例", "百分比，2位小数", "非成交或无值显示“-”"],
        ],
        weights=[1.25, 2.5, 1.6, 1.9],
        font_size=8.7,
    )

    add_heading(doc, "2.7 页面状态与异常提示", 2, page_break=True)
    add_table(
        doc,
        ["场景", "触发条件", "用户提示/展示", "可用操作"],
        [
            ["加载中", "首次进入、切换筛选或榜单", "骨架屏或加载指示，不闪烁旧数据", "允许返回"],
            ["空列表", "查询条件无匹配地块", "当前筛选条件下暂无匹配地块", "调整筛选、清空关键词"],
            ["榜单为空", "统计周期内无有效数据", "暂无榜单数据", "返回、切换榜单"],
            ["接口失败", "列表/详情/榜单请求失败", "数据加载失败，请稍后重试", "重试、返回"],
            ["网络异常", "离线或请求超时", "网络连接异常，请检查网络后重试", "重试、返回"],
            ["详情不存在", "地块ID失效或已下架", "该地块信息不存在或已下架", "返回土拍列表"],
            ["字段缺失", "单个非关键字段为空", "显示“-”，不隐藏整个模块", "无"],
            ["数据延迟", "数据尚未完成更新", "保留更新时间和免责声明", "查看官方公示"],
        ],
        weights=[1.2, 2.1, 2.8, 1.5],
        font_size=8.6,
    )
    add_callout(doc, "免责声明", "本数据由系统基于互联网公开信息采集并加工生成，非政府官方发布。数据仅供参考，可能存在延迟或偏差，不构成任何交易、投资或决策依据；具体情况以相关政府主管部门或官方机构最新公示为准。")

    add_heading(doc, "3. 数据与接口说明", 1)
    add_table(
        doc,
        ["模块", "建议接口", "关键入参", "关键返回", "失败处理"],
        [
            ["列表与统计", "GET /land-auctions", "keyword/status/district/usage/sort/page/pageSize", "状态数量、地块列表、起拍楼面价（元/㎡）、成交价（元/㎡）、发布时间、成交时间、分页信息", "展示错误态并可重试"],
            ["榜单", "GET /land-auctions/rankings", "type/cycle/limit", "榜单项、统计周期、更新时间", "榜单模块独立错误态"],
            ["地块详情", "GET /land-auctions/{id}", "id", "基础信息、交易信息、官方来源链接", "404展示下架态"],
            ["筛选字典", "GET /land-auctions/options", "city", "区域、用途、状态字典", "使用缓存字典或隐藏不可用条件"],
        ],
        weights=[1.2, 1.8, 2.2, 2.5, 1.6],
        font_size=8.2,
    )
    add_bullets(
        doc,
        [
            "正式环境不得使用前端推算成交总价、成交价或土地宗号；所有业务字段由后端返回并标识数据来源。",
            "状态数量应在同一查询口径下返回，避免Tab数量与列表结果不一致。",
            "接口需返回dataUpdatedAt，用于榜单更新时间和免责声明展示。",
            "建议保留officialSourceUrl、sourceName和sourcePublishedAt，便于纠错与追溯。",
        ],
    )

    add_heading(doc, "4. 埋点建议", 1)
    add_table(
        doc,
        ["事件", "触发时机", "建议属性"],
        [
            ["land_analysis_view", "进入土拍分析首页", "entry_source, city, user_login_status"],
            ["land_search", "关键词输入停止并触发查询", "keyword_length, result_count"],
            ["land_filter_apply", "确认筛选", "filter_type, filter_value, result_count"],
            ["land_status_switch", "切换状态Tab", "status, result_count"],
            ["land_ranking_switch", "切换榜单类型", "ranking_type, surface"],
            ["land_ranking_detail_view", "进入完整榜单", "ranking_type, cycle"],
            ["land_detail_view", "进入地块详情", "land_id, status, district, usage"],
            ["land_share_click", "点击分享", "land_id, share_channel"],
        ],
        weights=[2.0, 2.8, 3.7],
        font_size=8.6,
    )

    add_heading(doc, "5. 非功能要求", 1, page_break=True)
    add_table(
        doc,
        ["类别", "要求"],
        [
            ["兼容性", "适配主流iOS/Android微信内置浏览器；页面宽度320px及以上可用；支持安全区。"],
            ["性能", "首屏核心内容P75在3秒内可交互；筛选响应在500ms内给出反馈；图片与非首屏内容懒加载。"],
            ["可访问性", "按钮提供可读名称；键盘可操作；焦点可见；弹层打开后焦点进入弹层，关闭后返回触发点。"],
            ["安全与隐私", "仅展示公开土拍数据；接口参数校验并防注入；日志不记录用户完整搜索内容或敏感标识。"],
            ["数据准确性", "价格、面积、比例和日期必须保留数据源口径；禁止前端自行生成正式业务字段；展示更新时间与免责声明。"],
            ["可观测性", "记录接口成功率、P95耗时、空结果率、详情404率和数据更新时间延迟。"],
        ],
        weights=[1.3, 5.7],
        font_size=8.9,
    )

    add_heading(doc, "6. 验收标准", 1)
    add_table(
        doc,
        ["编号", "验收场景", "Given / When / Then"],
        [
            ["AC-01", "首页入口", "Given用户在首页，When点击土拍分析，Then进入土拍分析首页并显示默认榜单与地块列表。"],
            ["AC-02", "关键词搜索", "Given列表已加载，When输入地块名/编号/竞得公司关键词，Then列表仅展示包含匹配结果。"],
            ["AC-03", "组合筛选", "Given已选择状态，When再选择区域和用途并确认，Then按全部条件交集返回结果并回显条件。"],
            ["AC-04", "筛选取消", "Given筛选弹层已修改草稿值，When点击关闭/蒙层/Esc，Then列表和已生效条件保持不变。"],
            ["AC-05", "榜单切换", "Given位于首页或榜单详情，When切换榜单类型，Then标题、数值单位、列表内容同步更新。"],
            ["AC-06", "详情返回", "Given从列表某滚动位置进入详情，When返回，Then恢复原筛选条件及进入前滚动位置。"],
            ["AC-07", "异常兜底", "Given接口失败或无数据，When页面完成请求，Then显示对应错误/空状态且不出现脚本报错或空白页。"],
            ["AC-08", "列表时间与成交价", "Given列表加载完成，When地块状态为成交，Then以元/㎡展示起拍楼面价和成交价，并展示成交时间；When状态非成交，Then不展示成交价且成交时间显示“-”；所有地块均展示发布时间。"],
            ["AC-09", "免责声明", "Given用户浏览列表、榜单或详情，When查看页面底部，Then可看到数据来源边界与官方公示提示。"],
        ],
        weights=[0.9, 1.7, 5.2],
        font_size=8.6,
    )

    add_heading(doc, "7. 待确认项", 1)
    add_table(
        doc,
        ["问题", "影响", "建议"],
        [
            ["真实数据来源、采集合法性、更新频率和审核责任人是什么？", "数据准确性、合规与榜单可信度", "上线前明确官方来源、抓取/录入方式、审核SLA和纠错入口"],
            ["待出让/流拍/终止地块的详情摘要是否仍显示“成交楼面价/成交总价”？", "详情文案准确性", "按状态动态替换为起始价或显示“-”"],
            ["状态数量是否随区域、用途、关键词联动？", "筛选体验和接口设计", "建议随除状态外的其他条件联动更新"],
            ["榜单统计周期支持年度切换还是固定当年？", "榜单交互与接口参数", "若一期不支持切换，明确固定当年并显示完整年份"],
            ["更多操作和分享按钮的最终能力是什么？", "交互闭环", "明确分享卡片、复制链接、投诉纠错等范围；未实现前隐藏按钮"],
            ["是否需要B端土拍数据管理？", "真实数据上线和运营成本", "另立B端需求，覆盖采集、审核、发布、下架、纠错、字典与权限"],
        ],
        weights=[3.5, 2.0, 3.4],
        font_size=8.4,
    )

    # Document properties are intentionally generic and privacy-safe.
    props = doc.core_properties
    props.title = "成都住建房产超市土拍分析需求文档"
    props.subject = "土拍分析C端产品需求"
    props.author = "成都住建房产超市产品组"
    props.keywords = "土拍, 土地拍卖, 产品需求, PRD"
    props.comments = "V1.1 草案"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
